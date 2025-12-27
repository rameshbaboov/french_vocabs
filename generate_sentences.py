#!/usr/bin/env python3
import os, time, glob, argparse
from typing import List
import requests
from docx import Document

# -------- Defaults --------
INPUT_DIR = "out_french"          # folder with .txt word lists
OUTPUT_DIR = "out_sentences"      # folder for DOCX output
PROCESSED_FILE = "processed.txt"  # ledger to avoid reprocessing
OLLAMA_URL = "http://127.0.0.1:11434/api/generate"
MODEL_DEFAULT = "mistral:latest"
POLL_INTERVAL_DEFAULT = 5.0
TIMEOUT_DEFAULT = 300
RETRIES = 3
WORDS_PER_DOC = 10                # number of words per DOCX

# -------- Utils --------
def ensure_dir(p: str) -> None:
    os.makedirs(p, exist_ok=True)

def load_processed(path: str) -> set:
    if not os.path.exists(path):
        return set()
    return set(x.strip() for x in open(path, encoding="utf-8") if x.strip())

def append_processed(path: str, item: str) -> None:
    with open(path, "a", encoding="utf-8") as f:
        f.write(item + "\n")

def read_words(path: str) -> List[str]:
    return [ln.strip() for ln in open(path, encoding="utf-8") if ln.strip()]

def detect_level_from_filename(filename: str) -> str | None:
    name = os.path.basename(filename).upper()
    for lvl in ("A1", "A2", "B1", "B2"):
        if f"_{lvl}_" in name or name.startswith(f"FRENCH_{lvl}_"):
            return lvl
    return None

def post_ollama(model: str, prompt: str, system: str, timeout: int) -> str:
    payload = {
        "model": model,
        "prompt": prompt,
        "system": system,
        "options": {"temperature": 0.7, "top_p": 0.9},
        "stream": False
    }
    r = requests.post(OLLAMA_URL, json=payload, timeout=timeout)
    r.raise_for_status()
    data = r.json()
    if data.get("error"):
        raise RuntimeError(f"Ollama error: {data['error']}")
    resp = data.get("response", "")
    if not resp.strip():
        raise RuntimeError(f"Ollama empty response: {data}")
    return resp

# -------- Prompts --------
def build_prompt_meaning(word: str) -> tuple[str, str]:
    system = (
        "You are a translation assistant. "
        "Translate the given French word into English. "
        "Output only in the format: 'Word: <French word> | Meaning: <English>'."
    )
    prompt = f"Translate this French word into English: {word}"
    return system, prompt

def build_prompt_sentences(word: str, level: str) -> tuple[str, str]:
    system = (
        "You are a bilingual assistant. "
        "Create French learning material. "
        "No commentary, no notes, only the requested format."
    )
    prompt = (
        f"Niveau CECR : {level}\n\n"
        f"Word: {word}\n"
        f"Task: Write exactly 10 sentences using '{word}'.\n"
        "Each sentence must be B1 level French, followed immediately by its English translation.\n"
        "Format:\n"
        "1. <French sentence>\n"
        "   <English translation>\n"
        "2. <French sentence>\n"
        "   <English translation>\n"
        "... up to 10.\n"
        "- Never write 'Sentence 1:'. Only use the number and a period.\n"
    )
    return system, prompt

def write_docx(out_path: str, content: List[str]) -> None:
    doc = Document()
    for line in content:
        if line.strip() == "":
            doc.add_paragraph("")
        elif line.startswith("Word:"):
            para = doc.add_paragraph(line)
            para.runs[0].bold = True
        elif line.startswith("Meaning:") or "| Meaning:" in line:
            para = doc.add_paragraph(line)
            para.runs[0].italic = True
        else:
            doc.add_paragraph(line)
    doc.save(out_path)

# -------- Core --------
def process_file(in_file: str, out_root: str, model: str, default_level: str, timeout: int) -> bool:
    words = read_words(in_file)
    if not words:
        print(f"⚠️  {in_file}: no words")
        return True

    level = default_level or detect_level_from_filename(in_file)
    base = os.path.splitext(os.path.basename(in_file))[0]
    out_dir = os.path.join(out_root, base)
    ensure_dir(out_dir)

    print(f"▶ {in_file}: {len(words)} words → groups of {WORDS_PER_DOC} | level={level}")

    batches = [words[i:i + WORDS_PER_DOC] for i in range(0, len(words), WORDS_PER_DOC)]
    for b_idx, batch in enumerate(batches, start=1):
        all_content = []
        for word in batch:
            meaning_text = None
            sentences_text = None

            # Step A: Meaning
            for attempt in range(1, RETRIES+1):
                try:
                    print(f"⏳ Getting meaning for '{word}'...")
                    sys1, p1 = build_prompt_meaning(word)
                    meaning_text = post_ollama(model, p1, sys1, timeout).strip()
                    break
                except Exception as e:
                    print(f"⚠️ Meaning failed for '{word}' (try {attempt}): {e}")
                    time.sleep(2.0)
            if not meaning_text:
                meaning_text = f"Word: {word} | Meaning: ???"

            # Step B: Sentences
            for attempt in range(1, RETRIES+1):
                try:
                    print(f"⏳ Getting sentences for '{word}'...")
                    sys2, p2 = build_prompt_sentences(word, level)
                    sentences_text = post_ollama(model, p2, sys2, timeout).strip()
                    break
                except Exception as e:
                    print(f"⚠️ Sentences failed for '{word}' (try {attempt}): {e}")
                    time.sleep(2.0)
            if not sentences_text:
                sentences_text = "⚠️ No sentences generated."

            # Combine
            all_content.append(meaning_text)
            all_content.append("")
            all_content.extend([ln for ln in sentences_text.splitlines() if ln.strip()])
            all_content.append("")

        out_path = os.path.join(out_dir, f"{base}_part{b_idx:03d}.docx")
        write_docx(out_path, all_content)
        print(f"💾 Saved {out_path}")

    print(f"🎉 Done {in_file}")
    return True

# -------- Main --------
def main():
    ap = argparse.ArgumentParser(description="Generate bilingual French practice sentences in DOCX (two-step: meaning + sentences).")
    ap.add_argument("--input-dir", default=INPUT_DIR)
    ap.add_argument("--output-dir", default=OUTPUT_DIR)
    ap.add_argument("--model", default=MODEL_DEFAULT)
    ap.add_argument("--level", choices=["A1", "A2", "B1", "B2"], default="A1",
                    help="Default CEFR level (overrides filename detection).")
    ap.add_argument("--poll", type=float, default=POLL_INTERVAL_DEFAULT)
    ap.add_argument("--timeout", type=int, default=TIMEOUT_DEFAULT)
    args = ap.parse_args()

    ensure_dir(args.output_dir)
    processed_path = os.path.join(args.output_dir, PROCESSED_FILE)
    processed = load_processed(processed_path)

    print(f"👀 Watching {args.input_dir}/*.txt for new wordlists…")
    print(f"   Output → {args.output_dir}  |  Model → {args.model}")
    print(f"   Processed ledger: {processed_path}")
    print("Press Ctrl+C to stop.\n")

    try:
        while True:
            files = sorted(glob.glob(os.path.join(args.input_dir, "*.txt")))
            for f in files:
                if f in processed:
                    continue
                print(f"➡️ Found new file {f}")
                ok = process_file(
                    in_file=f,
                    out_root=args.output_dir,
                    model=args.model,
                    default_level=args.level,
                    timeout=args.timeout
                )
                if ok:
                    append_processed(processed_path, f)
                    processed.add(f)
            time.sleep(args.poll)
    except KeyboardInterrupt:
        print("👋 Stopped")

if __name__ == "__main__":
    main()
